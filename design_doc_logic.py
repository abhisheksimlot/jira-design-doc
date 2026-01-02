import os
import json
import io
from datetime import date
from io import BytesIO

from docx import Document
from openai import OpenAI
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Inches


def extract_json_from_text(text: str) -> str:
    """
    Extract the first JSON object found in the text (between the first '{' and last '}').
    Works even if the model adds extra words before/after JSON.
    """
    if not text:
        return ""

    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return ""

    return text[start:end + 1]


# Read API key from environment variable
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError(
        "OPENAI_API_KEY environment variable is not set. "
        "Please set it before running the app."
    )

# Model name – you can change to any deployed model you have
MODEL_NAME = "gpt-4.1-mini"

# Sections of your design document (adapt if you want)
SECTIONS = [
    "1. Overview",
    "1.1 Audience",
    "1.2 Scope",
    "1.3 Shared Responsibility Matrix",
    "1.4 Software Development Life Cycle (SDLC)",
    "2. Solution Overview",
    "2.1 Business Requirements",
    "2.2 Solution Summary",
    "2.3 Architectural Diagram",
    "2.4 Solution Components",
    "2.5 Custom APIs & Connectors",
    "2.6 Data Storage & Dataverse",
    "2.7 Integrations",
    "2.8 Automation Scheduling",
    "2.9 Exception Handling in Power Automate Flows",
    "2.10 Data & Retention Policy",
    "3. Security",
    "3.1 Authentication & Authorization",
    "3.2 App Security",
    "3.3 Security in DevOps",
    "4. Deployment & DevOps",
    "5. Non-Functional Requirements",
    "6. Risks & Assumptions",
    "7. References & Appendix",
]

client = OpenAI(api_key=OPENAI_API_KEY)

def generate_all_sections(jira_text: str) -> dict:
    system_prompt = (
        "You are a senior solution architect in an insurance company. "
        "You write clear, structured design documentation."
    )

    section_list = "\n".join([f"- {s}" for s in SECTIONS])

    user_prompt = f"""
You MUST return ONLY a valid JSON object (no markdown, no explanations, no headings).
The JSON must start with {{ and end with }}.

Rules:
1) Every key must be EXACTLY one of the section titles below.
2) Every value must be a string containing the section body text (do NOT include the title).
3) If you do not have enough information for a section, return an empty string for that section.

SECTION TITLES:
{section_list}

JIRA STORIES:
\"\"\"{jira_text}\"\"\"

Return ONLY JSON.
"""

    response = client.responses.create(
        model=MODEL_NAME,
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        max_output_tokens=4000,
    )

    # Older SDKs: get text like this
    raw_text = response.output[0].content[0].text
    # If your SDK already returns a string, this is fine.

    json_str = extract_json_from_text(raw_text)
    if not json_str:
        raise ValueError(
            "Model did not return valid JSON. Try reducing input text or try again."
        )

    return json.loads(json_str)



def generate_section_text(jira_text: str, section_title: str) -> str:
    """
    Use the LLM to generate ONE section of the design document
    from the full Jira stories text.
    """
    system_prompt = (
        "You are a senior solution architect in an insurance company. "
        "You write clear, structured design documentation for Power Platform "
        "solutions (Canvas Apps, Model-Driven Apps, Power Automate, Dataverse, "
        "custom APIs) aimed at architects, developers and business stakeholders."
    )

    user_prompt = f"""
Using the following Jira stories, write the section {section_title} for a
Solution Design Document.

- Keep it specific to the described solution.
- Use concise paragraphs and short bullet lists where helpful.
- Assume the platform is Microsoft Power Platform plus supporting services.
- Write in neutral, professional British English.
- Do NOT repeat the section title in the body; that will be added by the app.

JIRA STORIES:
\"\"\" 
{jira_text}
\"\"\"
"""

    response = client.responses.create(
        model=MODEL_NAME,
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        max_output_tokens=800,
    )

    return response.output[0].content[0].text.strip()


def generate_design_doc_bytes(
    jira_text: str,
    project_name: str,
    version: str,
    prepared_by: str,
) -> bytes:
    """
    Build the full design document in memory and return it as raw bytes.
    """

    # 🔹 Single OpenAI call to get ALL sections
    sections = generate_all_sections(jira_text)

    doc = Document()

    # Header
    doc.add_heading("Solution Design Document", level=0)
    doc.add_paragraph(f"Project Name: {project_name}")
    doc.add_paragraph(f"Version: {version}")
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    doc.add_paragraph(f"Prepared By: {prepared_by}")
    doc.add_page_break()

    # Sections from the JSON
    for section_title in SECTIONS:
        heading_level = 1 if section_title.count(".") == 1 else 2
        doc.add_heading(section_title, level=heading_level)

        if section_title == "2.3 Architectural Diagram":
            png_bytes = create_architecture_diagram_png()
            doc.add_paragraph("Diagram (auto-generated):")
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.5))
            doc.add_paragraph("")  # spacing

        body = sections.get(section_title, "")
        if not body:
            continue

        for para in body.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def create_architecture_diagram_png() -> bytes:
    """
    Creates a simple architecture diagram as a PNG (boxes + arrows) with BIGGER fonts.
    """
    # ✅ Bigger canvas + higher DPI feel
    W, H = 2200, 1200   # was 1400x800
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    # ✅ Bigger fonts
    try:
        font = ImageFont.truetype("arial.ttf", 34)     # was 22
        font_b = ImageFont.truetype("arial.ttf", 42)   # was 26
    except:
        font = ImageFont.load_default()
        font_b = ImageFont.load_default()

    def box(x1, y1, x2, y2, title, subtitle=""):
        draw.rounded_rectangle(
            [x1, y1, x2, y2],
            radius=26,                 # was 18
            outline="black",
            width=4,                   # was 3
            fill="#F5F9FF"
        )
        draw.text((x1 + 28, y1 + 22), title, fill="black", font=font_b)
        if subtitle:
            draw.text((x1 + 28, y1 + 85), subtitle, fill="black", font=font)

    def arrow(x1, y1, x2, y2):
        draw.line([x1, y1, x2, y2], fill="black", width=6)  # was 4

        # Bigger arrow head
        import math
        ang = math.atan2(y2 - y1, x2 - x1)
        L = 30  # was 18
        a1 = ang + math.radians(150)
        a2 = ang - math.radians(150)
        draw.line([x2, y2, x2 + L * math.cos(a1), y2 + L * math.sin(a1)], fill="black", width=6)
        draw.line([x2, y2, x2 + L * math.cos(a2), y2 + L * math.sin(a2)], fill="black", width=6)

    # ✅ Reposition boxes to fit the bigger canvas
    box(120, 120, 700, 300, "Business User", "Browser / Teams")
    box(820, 120, 1450, 300, "Power Apps", "Canvas / Model-driven")
    box(1550, 120, 2100, 300, "Power BI", "Embedded (optional)")

    box(820, 430, 1450, 630, "Power Automate", "Cloud flows / approvals")
    box(820, 760, 1450, 980, "Dataverse", "Core data store")

    box(120, 760, 700, 980, "SharePoint / Files", "Uploads / Templates")
    box(1550, 430, 2100, 630, "External Systems", "APIs / Mainframe / Email")
    box(1550, 760, 2100, 980, "Azure OpenAI", "Requirements / Summaries")

    # Arrows (adjusted)
    arrow(700, 210, 820, 210)        # user -> apps
    arrow(1135, 300, 1135, 430)      # apps -> flows
    arrow(1135, 630, 1135, 760)      # flows -> dataverse
    arrow(700, 870, 820, 870)        # sharepoint -> dataverse
    arrow(1450, 530, 1550, 530)      # flows -> external
    arrow(1450, 870, 1550, 870)      # dataverse -> openai
    arrow(1450, 210, 1550, 210)      # apps -> powerbi

    # Title (bigger)
    draw.text((120, 35), "High-Level Architecture", fill="black", font=font_b)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf.getvalue()


